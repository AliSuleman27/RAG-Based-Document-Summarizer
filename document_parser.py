
from abc import ABC, abstractmethod
from typing import Union
import os
from io import BytesIO

class DocumentParser(ABC):
    """Abstract base class for document parsers"""
    
    @abstractmethod
    def extract_text(self, file: Union[str, BytesIO]) -> str:
        """Extract text from a document file"""
        pass

    @classmethod
    def get_parser(cls, file: Union[str, BytesIO]):
        """Factory method to get the appropriate parser based on file extension"""
        if isinstance(file, str):
            filename = file
        else:
            filename = getattr(file, 'filename', '')
        
        if not filename:
            raise ValueError("Cannot determine file type - no filename provided")
        
        _, ext = os.path.splitext(filename)
        ext = ext.lower()
        
        if ext == '.pdf':
            return PDFParser()
        elif ext == '.txt':
            return TXTParser()
        elif ext in ('.doc', '.docx'):
            return DOCParser()
        else:
            raise ValueError(f"Unsupported file format: {ext}")

class PDFParser(DocumentParser):
    """Parser for PDF documents using PyPDF2"""
    
    def extract_text(self, file: Union[str, BytesIO]) -> str:
        from PyPDF2 import PdfReader
        
        if isinstance(file, str):
            with open(file, 'rb') as f:
                reader = PdfReader(f)
        else:
            reader = PdfReader(file)
            
        text = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
        return '\n'.join(text)

class TXTParser(DocumentParser):
    """Parser for plain text documents"""
    
    def extract_text(self, file: Union[str, BytesIO]) -> str:
        if isinstance(file, str):
            with open(file, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            file.seek(0)
            return file.read().decode('utf-8')

class DOCParser(DocumentParser):
    """Parser for Word documents using python-docx"""
    
    def extract_text(self, file: Union[str, BytesIO]) -> str:
        from docx import Document
        
        if isinstance(file, str):
            doc = Document(file)
        else:
            doc = Document(file)
            
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])